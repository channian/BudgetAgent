"""生成交接文件 .docx"""
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.top_margin    = Cm(2.0)
sec.bottom_margin = Cm(2.0)
sec.left_margin   = Cm(2.5)
sec.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
FONT = "Microsoft JhengHei"

def _set_run(run, size=10.5, bold=False, color=None, mono=False):
    run.font.name  = "Courier New" if mono else FONT
    run.font.size  = Pt(size)
    run.font.bold  = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if not mono:
        rPr = run._r.get_or_add_rPr()
        rf  = rPr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rPr.insert(0, rf)
        rf.set(qn("w:eastAsia"), FONT)

def _shade(para, fill="F2F3F4"):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)

def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        _set_run(run, size=15, bold=True, color=(44, 62, 80))
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        _set_run(run, size=12, bold=True, color=(52, 73, 94))
    return p

def body(text, bold=False):
    p = doc.add_paragraph(text)
    for run in p.runs:
        _set_run(run, bold=bold)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.3)
    for run in p.runs:
        _set_run(run)
    return p

def code(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Inches(0.4)
    for run in p.runs:
        _set_run(run, size=9, color=(30, 90, 150), mono=True)
    _shade(p)
    return p

def sp():
    doc.add_paragraph("")

# ═════════════════════════════════════════════════════════════════════════════
# TITLE
# ═════════════════════════════════════════════════════════════════════════════
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("AI Agent 預算審核平台\n交接文件")
_set_run(r, size=24, bold=True, color=(44, 62, 80))

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = s.add_run(f"產出日期：{datetime.date.today()}　　版本：v2.0")
_set_run(r2, size=11, color=(127, 140, 141))

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 1. 專案簡介
# ═════════════════════════════════════════════════════════════════════════════
h1("1. 專案簡介")
body("本系統為 ASE 智慧系統部門（Smart System Department）內部使用的 AI Agent 預算審核平台，"
     "透過「AI 自動初審（Layer 1）→ 專家人工複核（Layer 2）」兩層機制管理 AI Agent 專案預算的申請與審批流程。")
sp()

h2("1.1 核心功能")
for item in [
    "待簽核案件列表（篩選、排序、匯出 CSV/XLSX、匯入）",
    "預算單詳情頁（AI 審核結果 + 專家決策 + 時間軸）",
    "建立新預算單（4 個自由輸入欄位；AI 資料由 RPA 自動對應）",
    "AI Agent 圖書館（靜態展示頁）",
    "派發中心（設定 dispatch_date / budget_no，SLA 計時起點）",
    "權限管理中心（Admin Web UI 新增/編輯/重設密碼）",
    "側邊欄：依角色過濾選單 + 自助修改密碼",
    "通知鈴鐺（顯示真實未讀數量）",
]:
    bullet(item)
sp()

# ═════════════════════════════════════════════════════════════════════════════
# 2. 技術架構
# ═════════════════════════════════════════════════════════════════════════════
h1("2. 技術架構")

h2("2.1 整體架構")
body("Flask（Python 3.12）後端提供靜態前端頁面與 /api/ RESTful 路由；"
     "前端使用 React 18 UMD + Babel Standalone（無 bundler，<script type=\"text/babel\"> 直接載入 JSX）。")
sp()

h2("2.2 目錄結構")
for line in [
    "BudgetAgent/",
    "├── backend/",
    "│   ├── app.py              # Flask 入口，APScheduler SLA cron",
    "│   ├── config.py           # DB 連線設定、SECRET_KEY",
    "│   ├── db.py               # psycopg2 連線池 + cursor context manager",
    "│   ├── requirements.txt",
    "│   ├── routes/",
    "│   │   ├── auth.py         # 登入/登出/me/修改自己密碼",
    "│   │   ├── budgets.py      # 預算 CRUD、export、import、approve/reject",
    "│   │   ├── notifications.py",
    "│   │   └── users.py        # 使用者 CRUD（admin only）",
    "│   └── utils/",
    "│       ├── audit.py        # 寫入 audit_logs",
    "│       └── sla.py          # 72h SLA 提醒",
    "├── budget/                 # 前端（Flask 直接 serve）",
    "│   ├── index.html",
    "│   ├── app.jsx             # React 根元件、路由狀態機",
    "│   ├── api.js              # API client + DB ↔ 前端欄位對應",
    "│   ├── components.jsx      # Sidebar、共用元件",
    "│   ├── data.js             # 常數、MOCK 資料",
    "│   ├── page-list.jsx       # 列表頁",
    "│   ├── page-detail.jsx     # 詳情頁",
    "│   ├── page-edit.jsx       # 建立/編輯預算單",
    "│   └── page-other.jsx      # 派發中心、圖書館、權限管理",
    "├── rpa/",
    "│   └── ingest.py           # RPA JSON 批次匯入",
    "└── database/",
    "    ├── schema_v2.sql       # 完整 DB 建表語法（含 ALTER 升級區塊）",
    "    └── set_password.py     # CLI 工具：設定使用者初始密碼",
]:
    code(line)
sp()

# ═════════════════════════════════════════════════════════════════════════════
# 3. 資料庫
# ═════════════════════════════════════════════════════════════════════════════
h1("3. 資料庫")

h2("3.1 連線資訊（backend/config.py）")
for line in [
    'DB = {',
    '    "dbname":  "CIM",',
    '    "user":    "cim_admin",',
    '    "password": "1qaz2wsx3edc",',
    '    "host":    "10.10.51.98",   # 依實際內網 IP 修改',
    '    "port":    "5432",',
    '    "options": "-c search_path=budget",',
    '}',
    'SECRET_KEY = "pensieve-internal-2026"',
]:
    code(line)
sp()

h2("3.2 Schema 重要說明")
for item in [
    "資料庫：CIM　　Schema：budget",
    "status / expert_decision：VARCHAR + CHECK（非 ENUM），查詢時不加 ::text 轉型",
    "owner：VARCHAR 自由文字（非 INT FK → users），不要加 LEFT JOIN budget.users",
    "PK 欄位名稱統一是 id（SERIAL），舊文件出現的 jsondb_id 已廢棄",
    "ai_result：JSONB，psycopg2 插入須用 psycopg2.extras.Json() 包裝",
    "audit_logs.diff_before / diff_after：也是 JSONB（非 TEXT）",
]:
    bullet(item)
sp()

h2("3.3 budget.budget_requests（核心資料表）")
for col in [
    "id              SERIAL PK",
    "project_name    VARCHAR(300) UNIQUE NOT NULL  ← UPSERT 衝突鍵",
    "week            INT         ISO 週次，自動計算",
    "category        VARCHAR(100)",
    "sub_category    VARCHAR(100)",
    "expert_name     VARCHAR(100)",
    "owner           VARCHAR(100)  自由文字，非 FK",
    "amount          DECIMAL(15,2)",
    "ai_result       JSONB  {AI處置結果, 保留案件的信心分數}",
    "ai_comment      TEXT   RPA 原因",
    "expert_decision VARCHAR(10) CHECK IN ('通過','退件')",
    "expert_comment  TEXT",
    "dispatch_date   TIMESTAMP  SLA 計時起點",
    "sign_date       TIMESTAMP  專家完成簽核時間",
    "cycle_time      INT        sign_date - dispatch_date（天）",
    "status          VARCHAR(20) CHECK IN ('AI_REVIEW','EXPERT_REVIEW',",
    "                'PENDING_ACTION','CLOSED','REJECTED')",
    "note            TEXT",
    "created_at      TIMESTAMP DEFAULT NOW()",
]:
    bullet(col, level=1)
sp()

h2("3.4 budget.users")
for col in [
    "id          SERIAL PK",
    "name        VARCHAR",
    "department  VARCHAR",
    "ad_account  VARCHAR UNIQUE",
    "password    VARCHAR  werkzeug pbkdf2_sha256 hash",
    "role        VARCHAR CHECK IN ('admin','expert','viewer')",
    "email       VARCHAR",
]:
    bullet(col, level=1)
sp()

h2("3.5 Status 狀態流程")
body("AI_REVIEW → EXPERT_REVIEW → CLOSED")
body("                          ↘ PENDING_ACTION → EXPERT_REVIEW（退件重送）")
body("                          ↘ REJECTED")
sp()
body("Status 顏色對應（UI 需使用這些 hex 值）：", bold=True)
for s in [
    "AI_REVIEW      → #7c3aed（紫）",
    "EXPERT_REVIEW  → #06b6d4（青）",
    "PENDING_ACTION → #f59e0b（橙）",
    "CLOSED         → #10b981（綠）",
    "REJECTED       → #ef4444（紅）",
]:
    bullet(s)
sp()

# ═════════════════════════════════════════════════════════════════════════════
# 4. RBAC
# ═════════════════════════════════════════════════════════════════════════════
h1("4. 角色權限（RBAC）")

h2("4.1 角色 vs 功能模組")
tbl = doc.add_table(rows=4, cols=6)
tbl.style = "Table Grid"
headers = ["角色", "待簽核", "已簽核完成", "圖書館", "派發中心", "權限管理"]
for i, hdr in enumerate(headers):
    cell = tbl.cell(0, i)
    cell.text = hdr
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(10)

data = [
    ("admin",  "✔ 可編輯", "✔", "✔", "✔", "✔"),
    ("expert", "✔ 可審核", "✔", "✔", "✗", "✗"),
    ("viewer", "✔ 唯讀",   "✔", "✔", "✗", "✗"),
]
for r, row_data in enumerate(data, 1):
    for c, val in enumerate(row_data):
        tbl.cell(r, c).text = val
        for run in tbl.cell(r, c).paragraphs[0].runs:
            run.font.size = Pt(10)
sp()

h2("4.2 密碼管理")
for item in [
    "Admin 可在「權限管理中心」新增使用者並設定初始密碼",
    "Admin 可重設任何使用者的密碼（無法查看現有明文密碼，單向 hash）",
    "每位登入者均可在側邊欄頭像旁點擊「修改密碼」更換自己的密碼",
    "密碼使用 werkzeug generate_password_hash / check_password_hash（pbkdf2_sha256）",
]:
    bullet(item)
sp()

# ═════════════════════════════════════════════════════════════════════════════
# 5. RPA JSON 匯入
# ═════════════════════════════════════════════════════════════════════════════
h1("5. RPA JSON 匯入")

h2("5.1 JSON 格式（RPA 產出）")
for line in [
    '{',
    '  "案件名稱":  "ProjectX",',
    '  "判定類別":  "研發費用",',
    '  "判定系統":  "SubSys",',
    '  "負責專家":  "王小明",',
    '  "原因":     "符合預算規範",',
    '  "最終決策":  "通過",',
    '  "AI對於保留案件的信心分數": 85',
    '}',
]:
    code(line)
sp()

h2("5.2 欄位對應")
for mapping in [
    "案件名稱        → project_name（UPSERT 衝突鍵）",
    "判定類別        → category",
    "判定系統        → sub_category",
    "負責專家        → expert_name",
    "原因            → ai_comment",
    "最終決策 + 信心 → ai_result（JSONB）",
    "week            → 自動計算 datetime.now().isocalendar().week",
    "status          → 初始值 AI_REVIEW",
]:
    bullet(mapping)
sp()

h2("5.3 執行")
code("python rpa/ingest.py")
body("掃描 RPA 指定資料夾，逐一匯入 .json 檔，使用 UPSERT（project_name 為衝突鍵）。")
sp()

# ═════════════════════════════════════════════════════════════════════════════
# 6. 系統啟動
# ═════════════════════════════════════════════════════════════════════════════
h1("6. 系統啟動")

h2("6.1 環境需求")
for item in [
    "Python 3.10+，能連線至內網 PostgreSQL（10.10.51.98:5432）",
    "pip install -r backend/requirements.txt",
    "套件：Flask, psycopg2-binary, APScheduler, openpyxl",
]:
    bullet(item)
sp()

h2("6.2 啟動指令")
code("cd backend")
code("python app.py")
body("Flask 預設監聽 http://0.0.0.0:5000，前後端同源，直接瀏覽器開啟即可使用。")
sp()

h2("6.3 設定初始密碼（CLI）")
code("cd database")
code("python set_password.py")
body("依提示輸入 ad_account 與新密碼，即可為任一使用者設定初始登入密碼。")
sp()

h2("6.4 GitHub 儲存庫資訊")
body("Repository（GitHub）：", bold=True)
code("https://github.com/channian/BudgetAgent")
sp()
body("開發分支：", bold=True)
code("claude/elegant-dijkstra-fxlPR")
sp()
body("Clone 專案（首次）：")
code("git clone https://github.com/channian/BudgetAgent.git")
code("cd BudgetAgent")
code("git checkout claude/elegant-dijkstra-fxlPR")
sp()
body("更新至最新代碼（後續）：")
code("git fetch origin claude/elegant-dijkstra-fxlPR")
code("git checkout claude/elegant-dijkstra-fxlPR")
code("git pull origin claude/elegant-dijkstra-fxlPR")
sp()
body("推送修改：")
code("git add <changed_files>")
code('git commit -m "描述本次修改"')
code("git push -u origin claude/elegant-dijkstra-fxlPR")
sp()
body("注意：所有改動只推送到 claude/elegant-dijkstra-fxlPR，不要推送到 main / master。", bold=True)
sp()

# ═════════════════════════════════════════════════════════════════════════════
# 7. 已完成功能
# ═════════════════════════════════════════════════════════════════════════════
h1("7. 已完成功能清單")
done_items = [
    "Flask 後端骨架（Blueprint 分層：auth / budgets / users / notifications）",
    "psycopg2 連線池 + context manager（db.py）",
    "APScheduler SLA 72h cron 自動通知",
    "完整 DB schema（schema_v2.sql）含升級 ALTER 區塊",
    "登入 / 登出 / /api/auth/me（session-based）",
    "密碼 hash 驗證（werkzeug pbkdf2_sha256）",
    "自助修改密碼（PUT /api/auth/me/password）",
    "使用者 CRUD（admin only）：新增、編輯、重設密碼",
    "RBAC：admin / expert / viewer 三角色",
    "側邊欄依角色過濾選單、顯示真實使用者資訊",
    "建立預算單：4 欄自由輸入（project, category, owner, amount）",
    "AI 欄位自動從 DB 對應 project_name，不需人工填寫",
    "UPSERT：RPA 寫 AI 欄位、人工寫基本欄位，互不覆蓋",
    "審核詳情頁：專家決策、評論、時間軸",
    "匯出 CSV / XLSX（GET /api/budgets/export）",
    "匯入 CSV / XLSX（POST /api/budgets/import，彈性欄標對應）",
    "通知 API（/api/notifications）",
    "Audit log 寫入（每次狀態變更皆記錄 before/after JSONB）",
    "RPA ingest.py 批次匯入 JSON",
    "前端 RBAC：viewer 唯讀、expert 可審核、admin 全部功能",
    "權限管理中心：Admin Web UI 管理使用者",
    "set_password.py CLI 工具",
]
for item in done_items:
    bullet(item)
sp()

# ═════════════════════════════════════════════════════════════════════════════
# 8. 待完成功能
# ═════════════════════════════════════════════════════════════════════════════
h1("8. 待完成功能（移交事項）")

pending_items = [
    ("通知鈴鐺顯示真實未讀數量",
     "右上角鈴鐺目前是靜態 Badge，需接 GET /api/notifications 回傳 unread_count，建議每 60 秒輪詢"),
    ("類別篩選動態化",
     "列表頁類別下拉目前使用 data.js 常數，需改為從 DB 讀取實際出現的 category 值"),
    ("派發中心實作",
     "PUT /api/budgets/<id>：設定 dispatch_date 與 budget_no 欄位（SLA 計時起點），目前頁面為佔位符"),
    ("Expert Lock（並發鎖）",
     "同一張預算單同時只允許一名專家持有寫入權限，可使用 DB-level SELECT FOR UPDATE 或 Redis TTL 鎖"),
    ("列印審核單",
     "詳情頁「列印」按鈕 UI 已存在但無功能，需實作 window.print() 或後端 PDF 輸出（WeasyPrint / ReportLab）"),
    ("Windows AD 整合",
     "目前密碼由系統自管（werkzeug hash），未來可串接 LDAP3 讓使用者沿用 Windows 網域帳密登入"),
]
for title_txt, desc_txt in pending_items:
    p = doc.add_paragraph()
    r_t = p.add_run(f"◆ {title_txt}：")
    _set_run(r_t, bold=True)
    r_d = p.add_run(desc_txt)
    _set_run(r_d)
sp()

# ═════════════════════════════════════════════════════════════════════════════
# 9. 重要注意事項（給後續 AI）
# ═════════════════════════════════════════════════════════════════════════════
h1("9. 重要注意事項（給後續 AI）")

notes_items = [
    "不可修改 DB 資料表定義——schema 已在正式 DB 部署，只撰寫應用邏輯。",
    "status / expert_decision 是 VARCHAR CHECK，查詢 / 比較時不加 ::text 轉型。",
    "owner 是 VARCHAR 自由文字，不是 INT FK；不要加 LEFT JOIN budget.users。",
    "ai_result 是 JSONB，psycopg2 插入必須用 psycopg2.extras.Json() 包裝。",
    "PK 欄位名稱是 id（非 jsondb_id）；前端 api.js dbToFrontend() 使用 row.id。",
    "audit_logs.diff_before / diff_after 也是 JSONB（非 TEXT）。",
    "開發分支：claude/elegant-dijkstra-fxlPR（所有改動推到此分支）。",
    "前端無 bundler，JSX 以 <script type=\"text/babel\"> 執行，不可引入 Node.js/CommonJS 語法。",
    "雲端（Claude.ai Code）無法連線內網 DB（10.10.51.98），所有 DB 相關測試需在使用者 Windows PC 執行。",
    "密碼使用單向 hash，任何人（包括 Admin）均不可查看他人明文密碼，只能重設。",
    "created_at 欄位命名（非 updated_at），api.js 的 dbToFrontend 中 updatedAt 映射到 dispatch_date。",
]
for note in notes_items:
    bullet(note)
sp()

# ═════════════════════════════════════════════════════════════════════════════
# Footer
# ═════════════════════════════════════════════════════════════════════════════
doc.add_paragraph("─" * 60)
ft = doc.add_paragraph(f"文件產出日期：{datetime.date.today()}　　AI Agent 預算審核平台 v2.0")
ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in ft.runs:
    _set_run(run, size=9, color=(127, 140, 141))

# ── Save ──────────────────────────────────────────────────────────────────────
out = "/home/user/BudgetAgent/database/交接文件_BudgetAgent_v2.docx"
doc.save(out)
print(f"Saved: {out}")
